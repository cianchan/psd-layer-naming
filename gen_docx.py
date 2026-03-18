#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "PSD图层重命名工具说明.docx")

doc = Document()

# ── Page setup: A4, normal margins ──────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ── Helper: set font on a run ────────────────────────────────────────────────
def set_run_font(run, size=None, bold=False, color=None):
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    # font name (east-asian + ascii)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'PingFang SC')
    rFonts.set(qn('w:hAnsi'),   'PingFang SC')
    rFonts.set(qn('w:eastAsia'),'PingFang SC')
    rFonts.set(qn('w:cs'),      'PingFang SC')
    rPr.insert(0, rFonts)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, size=(16 if level == 1 else 13), bold=True,
                     color=((31, 73, 125) if level == 1 else (47, 84, 150)))
    return p

def add_paragraph(doc, text='', size=11):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size)
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    return p

def add_numbered_list(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        # bold lead + normal text
        if '**' in item:
            parts = item.split('**')
            # parts: ['', 'bold', ' rest']
            for j, part in enumerate(parts):
                if part == '':
                    continue
                run = p.add_run(part)
                set_run_font(run, size=11, bold=(j % 2 == 1))
        else:
            run = p.add_run(item)
            set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)

def set_cell_bg(cell, hex_color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    """Add borders to all cells in the table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_code_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('PSD 图层自动重命名工具 — 说明文档')
set_run_font(run, size=20, bold=True, color=(31, 73, 125))
title_p.paragraph_format.space_after = Pt(12)

doc.add_paragraph()  # spacer

# ════════════════════════════════════════════════════════════════════════════
# 一、工具背景与目的
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '一、工具背景与目的', level=1)
add_paragraph(doc,
    '设计团队在制作电商产品主图时，每个 PSD 文件的图层结构各不相同，图层名称随意。'
    '为了让后续自动化程序（如广告投放、模板系统）能稳定识别图层，需要将 PSD 图层统一命名成规范格式。')
add_paragraph(doc,
    '本工具通过 Flask 本地 Web 界面 + ExtendScript 控制 Adobe Photoshop，'
    '实现批量处理 PSD 文件的全流程自动化。')

# ════════════════════════════════════════════════════════════════════════════
# 二、运行环境要求
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '二、运行环境要求', level=1)
env_items = [
    '系统：macOS 12 Monterey 或以上',
    '软件：Adobe Photoshop 2023 / 2024 / 2025（已激活并安装在本机）',
    '语言：Python 3.9+',
    '依赖包：flask、python-docx',
    '访问地址：http://127.0.0.1:7861（本机浏览器打开）',
]
add_bullet_list(doc, env_items)

# ════════════════════════════════════════════════════════════════════════════
# 三、处理流程
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '三、处理流程（按顺序执行）', level=1)

steps = [
    ('收集隐藏图层 ID',
     '在做任何修改之前，先扫描全部图层，记录所有不可见图层的 ID（包括 visible=false、opacity=0 的图层）。'
     '此步骤保留原始隐藏状态，防止后续解锁操作改变可见性判断。'),
    ('解锁所有图层',
     '对文档内所有图层（含图层组）移除锁定，避免后续删除/重命名报错 8800。'),
    ('删除隐藏图层',
     '根据第 1 步记录的 ID，删除全部隐藏图层。同时删除完全空白（无任何像素内容，bounds 宽高为 0）的图层。'),
    ('应用图层蒙版',
     '对像素图层（NORMAL 类型）执行 Apply Mask，将蒙版合并入图层像素。非像素图层（如 SOLIDFILL）的蒙版在后续栅格化时处理。'),
    ('解散所有图层组（Ungroup）',
     '递归将所有 Group 组内的图层移出，恢复为平铺结构，保持图层原有顺序，最后删除空组。'),
    ('跳过无文本图层的 PSD',
     '若文档中不存在任何文本图层，则跳过该文件，不输出到目标文件夹。'),
    ('统计图层数量',
     '按类型统计剩余各类图层数量，用于判断是否需要编号。'),
    ('重命名 + 栅格化',
     '按命名规则重命名每个图层，同时对非文本图层执行栅格化（转为像素层），并在栅格化后应用剩余蒙版。'),
    ('保存输出',
     '将处理后的 PSD 文件另存为到指定输出文件夹，文件名与原文件相同，不覆盖原始文件。'),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    run_title = p.add_run(title + '  ')
    set_run_font(run_title, size=11, bold=True)
    run_desc = p.add_run(desc)
    set_run_font(run_desc, size=11)
    p.paragraph_format.space_after = Pt(5)

# ════════════════════════════════════════════════════════════════════════════
# 四、图层命名规则（Table）
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '四、图层命名规则', level=1)

headers = ['图层类型', '命名结果', '说明']
rows = [
    ('文本图层（Text）',
     'text / text1 / text2 …',
     '只有一个文本层时命名为 text，多个时按顺序编号'),
    ('像素图层（Pixel）/ 智能对象（Smart Object）',
     'scenebg',
     '统一命名 scenebg，不编号'),
    ('形状图层 / 填充图层\n（Shape / SolidFill / GradientFill）',
     'stickerbg / stickerbg1 / stickerbg2 …',
     '多个时按顺序编号'),
    ('边框图层（Frame）',
     'frame',
     '自动识别：形状/填充类图层，且宽高各覆盖画布 ≥ 80% 时判定为边框图层'),
]

table = doc.add_table(rows=1 + len(rows), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_borders(table)

# Header row
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_bg(hdr_cells[i], 'BDD7EE')   # light blue header
    for para in hdr_cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            set_run_font(run, size=11, bold=True)

# Data rows
for r_idx, row_data in enumerate(rows):
    row_cells = table.rows[r_idx + 1].cells
    for c_idx, val in enumerate(row_data):
        row_cells[c_idx].text = val
        if r_idx % 2 == 1:
            set_cell_bg(row_cells[c_idx], 'DEEAF1')  # alternate row tint
        for para in row_cells[c_idx].paragraphs:
            for run in para.runs:
                set_run_font(run, size=11)

# column widths
table.columns[0].width = Cm(5.5)
table.columns[1].width = Cm(5.0)
table.columns[2].width = Cm(7.5)

doc.add_paragraph()  # spacer after table

# ════════════════════════════════════════════════════════════════════════════
# 五、自动识别边框图层逻辑
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '五、自动识别边框图层（Frame Detection）逻辑', level=1)
add_paragraph(doc, '判断一个图层是否为"边框图层"的条件（同时满足）：')

frame_conds = [
    '图层类型为 非文本、非智能对象、非像素图层（即形状类/填充类图层）',
    '图层的宽度 ≥ 画布宽度 × 80%',
    '图层的高度 ≥ 画布高度 × 80%',
]
add_numbered_list(doc, frame_conds)

add_paragraph(doc, '满足以上三个条件则命名为 frame，否则按形状图层处理命名为 stickerbg。')

# ════════════════════════════════════════════════════════════════════════════
# 六、自动删除的图层类型
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '六、自动删除的图层类型', level=1)
add_paragraph(doc, '以下图层会在处理过程中自动删除：')
del_items = [
    '隐藏图层（小眼睛关闭 / opacity=0 / fillOpacity=0）',
    '完全空白像素图层（无任何像素内容，bounds 宽高均为 0）',
]
add_bullet_list(doc, del_items)

# ════════════════════════════════════════════════════════════════════════════
# 七、跳过条件
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '七、跳过条件', level=1)
add_paragraph(doc,
    '若 PSD 文件在删除隐藏图层后不包含任何文本图层，则该文件不进行任何重命名，'
    '也不输出到目标文件夹。日志中显示：')
add_code_para(doc, 'Skipped (no text layers): 文件名')

# ════════════════════════════════════════════════════════════════════════════
# 八、界面操作说明
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '八、界面操作说明', level=1)

ui_steps = [
    '打开浏览器访问 http://127.0.0.1:7861',
    '点击「选择输入文件夹」，选择包含原始 PSD 的文件夹',
    '点击「选择输出文件夹」，选择处理后文件的保存位置',
    '点击「开始处理」',
    '界面实时显示处理进度和日志',
    '处理完成后，输出文件夹中出现命名规范的 PSD 文件',
]
add_numbered_list(doc, ui_steps)

add_heading(doc, '注意事项', level=2)
notes = [
    '原始 PSD 文件不会被覆盖',
    '处理过程中 Photoshop 会自动打开（后台运行），处理完成后自动关闭文档',
    '如需中途停止，点击「停止」按钮',
]
add_bullet_list(doc, notes)

# ════════════════════════════════════════════════════════════════════════════
# 九、服务管理
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '九、服务管理', level=1)
add_paragraph(doc,
    '工具通过 macOS launchd 服务管理，登录时自动启动，崩溃自动重启。')

svc_items = [
    '日志路径：~/Library/Logs/psd-layer-naming.log',
    '停止服务：launchctl unload ~/Library/LaunchAgents/com.psd-layer-naming.plist',
    '启动服务：launchctl load ~/Library/LaunchAgents/com.psd-layer-naming.plist',
]
add_bullet_list(doc, svc_items)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f'Saved: {OUTPUT_PATH}')
