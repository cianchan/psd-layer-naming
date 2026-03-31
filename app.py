#!/usr/bin/env python3
"""
PSD Layer Renaming Tool
Reads layer naming rules from a .docx file,
generates an ExtendScript (.jsx) that Photoshop runs to:
  - Rename layers per naming rules
  - Delete hidden layers
  - Unlock locked layers
  - Rasterize smart objects and shape layers
  - Keep text layers as text (no rasterize)
  - Save renamed PSD to a new output folder
"""
import os
import re
import json
import glob
import threading
import time
import subprocess
from pathlib import Path
from flask import Flask, render_template, request, jsonify
from docx import Document

def find_photoshop_app_name():
    """Return the exact installed Photoshop app name for use in AppleScript tell blocks."""
    try:
        result = subprocess.run(
            ["mdfind", "kMDItemCFBundleIdentifier == 'com.adobe.Photoshop'"],
            capture_output=True, text=True, timeout=10
        )
        for path in result.stdout.strip().splitlines():
            if path.endswith(".app"):
                name = os.path.basename(path).replace(".app", "")
                if "Photoshop" in name:
                    return name
    except Exception:
        pass
    return "Adobe Photoshop"

app = Flask(__name__)
# ── Global job state ──────────────────────────────────────────
job_state = {
    "running": False,
    "logs": [],
    "progress": 0,
    "total": 0,
    "current_file": ""
}
ps_process = None  # holds the running osascript Popen handle
def log(msg):
    ts = time.strftime("%H:%M:%S")
    entry = f"[{ts}] {msg}"
    job_state["logs"].append(entry)
    print(entry)
# ── Routes ────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")
@app.route("/api/parse_doc", methods=["POST"])
def parse_doc():
    """Parse .docx and return extracted naming rules for preview."""
    data = request.json
    doc_path = data.get("doc_path", "").strip()
    if not doc_path or not os.path.exists(doc_path):
        return jsonify({"error": "文档路径无效"}), 400
    try:
        rules = extract_naming_rules(doc_path)
        return jsonify({"rules": rules})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/api/preview_psd", methods=["POST"])
def preview_psd():
    """List PSD files in the input folder."""
    data = request.json
    folder = data.get("input_folder", "").strip()
    if not folder or not os.path.isdir(folder):
        return jsonify({"error": "文件夹路径无效"}), 400
    psds = [Path(p).name for p in glob.glob(os.path.join(folder, "**/*.psd"), recursive=True)]
    psds += [Path(p).name for p in glob.glob(os.path.join(folder, "**/*.PSD"), recursive=True)]
    psds = sorted(set(psds))
    return jsonify({"files": psds, "count": len(psds)})
@app.route("/api/generate", methods=["POST"])
def generate():
    """Generate the ExtendScript .jsx file and optionally launch Photoshop."""
    data = request.json
    doc_path     = data.get("doc_path", "").strip()
    input_folder = data.get("input_folder", "").strip()
    output_folder= data.get("output_folder", "").strip()
    jsx_output   = data.get("jsx_output", "").strip() or os.path.join(os.path.expanduser("~"), "Desktop", "psd_rename.jsx")
    custom_rules = data.get("custom_rules", [])   # [{keyword, layer_name}, ...]
    for label, val in [("Doc路径", doc_path), ("输入文件夹", input_folder), ("输出文件夹", output_folder)]:
        if not val:
            return jsonify({"error": f"请填写 {label}"}), 400
    try:
        rules = custom_rules if custom_rules else extract_naming_rules(doc_path)
        jsx = build_jsx(input_folder, output_folder, rules)
        Path(jsx_output).write_text(jsx, encoding="utf-8")
        log(f"✅ JSX 脚本已生成: {jsx_output}")
        return jsonify({
            "ok": True,
            "jsx_path": jsx_output,
            "rule_count": len(rules),
            "jsx_preview": jsx[:800]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/api/pick_folder", methods=["POST"])
def pick_folder():
    """Open a native macOS folder picker and return the selected path."""
    data = request.json or {}
    prompt = data.get("prompt", "选择文件夹")
    try:
        result = subprocess.run(
            ["osascript", "-e", f'POSIX path of (choose folder with prompt "{prompt}")'],
            capture_output=True, text=True, timeout=60
        )
        path = result.stdout.strip()
        if not path:
            return jsonify({"cancelled": True})
        return jsonify({"path": path})
    except subprocess.TimeoutExpired:
        return jsonify({"cancelled": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/pick_file", methods=["POST"])
def pick_file():
    """Open a native macOS file picker for .docx files."""
    try:
        script = (
            'POSIX path of (choose file with prompt "选择命名规则文档" '
            'of type {"docx", "com.microsoft.word.doc"})'
        )
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=60
        )
        path = result.stdout.strip()
        if not path:
            return jsonify({"cancelled": True})
        return jsonify({"path": path})
    except subprocess.TimeoutExpired:
        return jsonify({"cancelled": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/run", methods=["POST"])
def run():
    """Generate JSX (type-based naming) and launch Photoshop automatically."""
    global ps_process
    data = request.json
    input_folder  = data.get("input_folder", "").strip()
    output_folder = data.get("output_folder", "").strip()
    # type_rules: {pixel, smart, shape} → name strings
    type_rules = data.get("type_rules", {
        "pixel": "scenebg", "smart": "scenebg", "shape": "stickerbg"
    })

    for label, val in [("输入文件夹", input_folder), ("输出文件夹", output_folder)]:
        if not val:
            return jsonify({"error": f"请填写 {label}"}), 400

    jsx_log_path = os.path.join(output_folder, "psd_renamer_jsx.log")
    stop_flag_path = os.path.join(output_folder, "psd_renamer_stop.flag")
    Path(stop_flag_path).unlink(missing_ok=True)   # clear any residual stop flag
    job_state["stop_flag_path"] = stop_flag_path
    try:
        jsx = build_jsx(input_folder, output_folder, type_rules, jsx_log_path, stop_flag_path)
        jsx_path = os.path.join(os.path.expanduser("~"), ".psd_renamer_tmp.jsx")
        Path(jsx_path).write_text(jsx, encoding="utf-8")
        Path(jsx_log_path).write_text("", encoding="utf-8")  # clear previous log
        log(f"📄 JSX: {jsx_path}")

        # Pass a short JS loader string — reads and evals the JSX file from within Photoshop.
        # Must use the EXACT installed app name so osascript loads Photoshop's dictionary;
        # without the dictionary, "do javascript" is unrecognized and causes a parse error.
        jsx_loader = (
            f"var _f=new File('{jsx_path}');"
            "_f.open('r');var _s=_f.read();_f.close();eval(_s);"
        )
        ps_app_name = find_photoshop_app_name()
        log(f"🎯 Photoshop 应用名: {ps_app_name}")
        ascript_path = os.path.join(os.path.expanduser("~"), ".psd_renamer_tmp.applescript")
        apple_script = (
            f'tell application "{ps_app_name}"\n'
            f'    activate\n'
            f'    delay 2\n'
            f'    do javascript "{jsx_loader}"\n'
            f'end tell'
        )
        Path(ascript_path).write_text(apple_script, encoding="utf-8")

        ps_process = subprocess.Popen(["osascript", ascript_path], stderr=subprocess.PIPE)
        job_state["running"] = True
        log("🚀 Photoshop 已启动，正在执行脚本…")

        def _watch(proc):
            # Poll log file for real-time updates while script runs
            last_size = 0
            while proc.poll() is None:
                time.sleep(2)
                try:
                    content = Path(jsx_log_path).read_text(encoding="utf-8")
                    if len(content) > last_size:
                        for line in content[last_size:].splitlines():
                            if line.strip():
                                log(f"[PS] {line}")
                        last_size = len(content)
                except Exception:
                    pass
            # Flush remaining log lines after process exits
            _, stderr = proc.communicate()
            try:
                content = Path(jsx_log_path).read_text(encoding="utf-8")
                if len(content) > last_size:
                    for line in content[last_size:].splitlines():
                        if line.strip():
                            log(f"[PS] {line}")
            except Exception:
                pass
            if stderr:
                err = stderr.decode().strip()
                if err:
                    log(f"❌ 错误: {err}")
            job_state["running"] = False
            log("✅ Photoshop 处理完成")
        threading.Thread(target=_watch, args=(ps_process,), daemon=True).start()

        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/stop", methods=["POST"])
def stop():
    """Kill the running Photoshop/osascript process."""
    global ps_process
    # Write stop flag so JSX stops before the next PSD
    if job_state.get("stop_flag_path"):
        try:
            Path(job_state["stop_flag_path"]).touch()
        except Exception:
            pass
    if ps_process and ps_process.poll() is None:
        ps_process.kill()
    ps_process = None
    job_state["running"] = False
    log("⛔ 已终止处理")
    return jsonify({"ok": True})

@app.route("/api/status")
def status():
    return jsonify({
        "running": job_state["running"],
        "logs": job_state["logs"][-200:],
        "progress": job_state["progress"],
        "total": job_state["total"],
        "current_file": job_state["current_file"]
    })
# ── Doc Parser ────────────────────────────────────────────────
def extract_naming_rules(doc_path):
    """
    Extract layer naming rules from the Word doc.
    Returns a list of dicts: {keyword, layer_name, description}
    """
    doc = Document(doc_path)
    rules = []
    # Pull all text from tables (the naming table in the doc)
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        # Find header row: look for a row containing "Type" or "Layer Naming"
        header_idx = None
        for i, row in enumerate(rows):
            combined = " ".join(row).lower()
            if "type" in combined or "layer naming" in combined or "命名" in combined:
                header_idx = i
                break
        if header_idx is None:
            continue
        headers = rows[header_idx]
        # Find columns
        type_col = next((i for i,h in enumerate(headers) if "type" in h.lower() or "类型" in h), None)
        name_col = next((i for i,h in enumerate(headers) if "naming" in h.lower() or "命名" in h.lower()), None)
        desc_col = next((i for i,h in enumerate(headers) if "display" in h.lower() or "说明" in h.lower() or "描述" in h.lower()), None)
        for row in rows[header_idx+1:]:
            if not any(row):
                continue
            rule = {
                "keyword": row[type_col].strip() if type_col is not None and type_col < len(row) else "",
                "layer_name": row[name_col].strip() if name_col is not None and name_col < len(row) else "",
                "description": row[desc_col].strip() if desc_col is not None and desc_col < len(row) else ""
            }
            # Clean: remove spaces from layer names (doc rule: no spaces)
            rule["layer_name"] = rule["layer_name"].replace(" ", "").replace("、", "/")
            if rule["keyword"] or rule["layer_name"]:
                rules.append(rule)
    # Fallback: scan paragraphs for patterns like "xxx命名" or "layer_name: xxx"
    if not rules:
        for para in doc.paragraphs:
            text = para.text.strip()
            # Match patterns like "product", "text1", "icon1" etc.
            m = re.search(r'([a-zA-Z][a-zA-Z0-9]*)[\s：:]+(.+)', text)
            if m:
                kw = m.group(1).lower()
                if kw in ("product","text","icon","logo","frame","scene","sticker","scenebg","stickerbg","bg"):
                    rules.append({"keyword": kw, "layer_name": kw, "description": text})
    return rules
# ── JSX Builder ───────────────────────────────────────────────
def build_jsx(input_folder, output_folder, type_rules, jsx_log_path="/tmp/psd_renamer_jsx.log", stop_flag_path="/tmp/psd_renamer_stop.flag"):
    """
    Build a Photoshop ExtendScript (.jsx) that renames layers purely by type.
    type_rules: dict with keys pixel, smart, shape -> target name strings.
    - Text layers  -> text1, text2, text3 ... (always numbered)
    - Pixel/Smart  -> e.g. "product"  (no number if only 1, else product1 product2...)
    - Shape        -> e.g. "stickerbg"
    - Hidden layers are deleted; smart objects & shapes are rasterized.
    """
    name_pixel = type_rules.get("pixel", "scenebg")
    name_smart = type_rules.get("smart", "scenebg")
    name_shape = type_rules.get("shape", "stickerbg")
    name_frame = type_rules.get("frame", "frame")
    input_folder_js  = input_folder.replace("\\", "/")
    output_folder_js = output_folder.replace("\\", "/")
    log_path_js = jsx_log_path.replace("\\", "/")
    stop_flag_path_js = stop_flag_path.replace("\\", "/")
    jsx = f"""// Auto-generated by PSD Layer Renaming Tool
#target photoshop
app.displayDialogs = DialogModes.NO;
var NAME_TEXT  = "text";
var NAME_PIXEL = "{name_pixel}";
var NAME_SMART = "{name_smart}";
var NAME_SHAPE = "{name_shape}";
var NAME_FRAME = "{name_frame}";
var g_workDoc   = null;  // set per-PSD in processPSD before countVisible
var g_frameCache = {{}};  // layerId → bool, cache per PSD to avoid duplicate detection work
var INPUT_FOLDER  = new Folder("{input_folder_js}");
var OUTPUT_FOLDER = new Folder("{output_folder_js}");
var LOG_FILE = new File("{log_path_js}");
var STOP_FILE = new File("{stop_flag_path_js}");
function jsxLog(msg) {{
    LOG_FILE.open("a");
    LOG_FILE.writeln(msg);
    LOG_FILE.close();
}}
// Numeric LayerKind values — PS 2025 deprecated some enum names; use numbers directly
var K_TEXT   = 2;   // LayerKind.TEXT
var K_NORMAL = 1;   // LayerKind.NORMAL
var K_SMART  = 17;  // LayerKind.SMARTOBJECT
var K_SHAPE  = 7;   // LayerKind.SHAPE
try {{ K_TEXT   = LayerKind.TEXT;        }} catch(e) {{}}
try {{ K_NORMAL = LayerKind.NORMAL;      }} catch(e) {{}}
try {{ K_SMART  = LayerKind.SMARTOBJECT; }} catch(e) {{}}
try {{ K_SHAPE  = LayerKind.SHAPE;       }} catch(e) {{}}
jsxLog("Script started");
jsxLog("Input: " + INPUT_FOLDER.fullName);
jsxLog("Input exists: " + INPUT_FOLDER.exists);
jsxLog("Output: " + OUTPUT_FOLDER.fullName);
if (!OUTPUT_FOLDER.exists) {{
    OUTPUT_FOLDER.create();
    jsxLog("Created output folder");
}}
// --- Visual frame detection: shape/fill layer that wraps all 4 edges of the canvas ---
// A "frame" layer is any non-text / non-smart-object / non-pixel layer whose bounding
// box covers ≥ 80% of both the document width AND height (e.g. a green rectangle border).
// This avoids the false positives that pixel-manipulation approaches produce on smart objects
// and the false negatives caused by SOLIDFILL layers that fill their entire rectangle.
function isFrameStyleLayer(layer) {{
    if (!g_workDoc) return false;
    if (g_frameCache.hasOwnProperty(layer.id)) return g_frameCache[layer.id];
    var result = false;
    try {{
        // Frame candidates are shape/fill layers only — skip text, smart objects, pixel layers
        var k = null;
        try {{ k = layer.kind; }} catch(e) {{}}
        var ks = "" + k;
        var isText   = (k === K_TEXT   || k === 2  || ks === "LayerKind.TEXT");
        var isSmart  = (k === K_SMART  || k === 17 || ks === "LayerKind.SMARTOBJECT");
        var isPixel  = (k === K_NORMAL || k === 1  || ks === "LayerKind.NORMAL");
        if (isText || isSmart || isPixel) {{ g_frameCache[layer.id] = false; return false; }}
        // The layer's bounds must cover ≥ 80% of both document dimensions
        var bounds = layer.bounds;
        var lx = +bounds[0], ty = +bounds[1], rx = +bounds[2], by = +bounds[3];
        var w = rx - lx, h = by - ty;
        var docW = +g_workDoc.width, docH = +g_workDoc.height;
        result = (w >= docW * 0.8 && h >= docH * 0.8);
        jsxLog("frame check [" + layer.name + "] w=" + Math.round(w) + "/" + Math.round(docW) +
               " h=" + Math.round(h) + "/" + Math.round(docH) + " → " + result);
    }} catch(e) {{
        jsxLog("frame check err [" + layer.name + "]: " + e.message);
        result = false;
    }}
    g_frameCache[layer.id] = result;
    return result;
}}
// --- Get base name by layer type (visual frame check takes priority over type) ---
function getBaseName(layer) {{
    var k = null;
    try {{ k = layer.kind; }} catch(e) {{}}
    var ks = "" + k;
    // Text layers are always "text"
    if (k === K_TEXT   || k === 2  || ks === "LayerKind.TEXT")        return NAME_TEXT;
    // Visual frame detection: hollow-center layers → "frame"
    if (isFrameStyleLayer(layer))                                      return NAME_FRAME;
    if (k === K_SMART  || k === 17 || ks === "LayerKind.SMARTOBJECT") return NAME_SMART;
    if (k === K_NORMAL || k === 1  || ks === "LayerKind.NORMAL")      return NAME_PIXEL;
    // SOLIDFILL(3), GRADIENTFILL(4), PATTERNFILL(5), SHAPE(7), or unknown → shape
    return NAME_SHAPE;
}}
// --- Get ACTUAL displayed visibility via Action Manager (includes Layer Comp state) ---
function getAMVisible(layer) {{
    try {{
        var ref = new ActionReference();
        ref.putIdentifier(charIDToTypeID("Lyr "), layer.id);
        var desc = executeActionGet(ref);
        return desc.getBoolean(stringIDToTypeID("visible"));
    }} catch(e) {{
        return layer.visible;
    }}
}}
// --- Comprehensive hidden check: any one of these being true means the layer is hidden ---
function isHidden(layer) {{
    try {{ if (!layer.visible)     return true; }} catch(e) {{}}
    try {{ if (!getAMVisible(layer)) return true; }} catch(e) {{}}
    try {{ if (layer.opacity    === 0) return true; }} catch(e) {{}}
    try {{ if (layer.fillOpacity === 0) return true; }} catch(e) {{}}
    return false;
}}
// --- Collect hidden layer IDs BEFORE any modifications (preserves original state) ---
function collectHiddenIds(layerSet, hiddenIds) {{
    for (var i = 0; i < layerSet.artLayers.length; i++) {{
        try {{
            var l = layerSet.artLayers[i];
            var propVis = l.visible;
            var amVis   = getAMVisible(l);
            var opac    = l.opacity;
            var fillOp  = l.fillOpacity;
            var hidden  = !propVis || !amVis || opac === 0 || fillOp === 0;
            jsxLog("check [" + l.name + "] prop=" + propVis + " am=" + amVis +
                   " opacity=" + opac + " fillOpacity=" + fillOp + " hidden=" + hidden);
            if (hidden) {{ hiddenIds[l.id] = true; continue; }}
            // Also delete completely empty pixel layers (bounds width+height = 0)
            try {{
                var lk = null;
                try {{ lk = l.kind; }} catch(e) {{}}
                var lks = "" + lk;
                var isText = (lk === K_TEXT || lk === 2 || lks === "LayerKind.TEXT");
                if (!isText) {{
                    var b = l.bounds;
                    var w = b[2] - b[0];
                    var h = b[3] - b[1];
                    if (w == 0 && h == 0) {{
                        jsxLog("check [" + l.name + "] empty (no pixels) → delete");
                        hiddenIds[l.id] = true;
                    }}
                }}
            }} catch(e) {{}}
        }} catch(e) {{ jsxLog("collectHidden err: " + e.message); }}
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        try {{
            var grp = layerSet.layerSets[j];
            var grpHidden = !grp.visible || !getAMVisible(grp);
            if (grpHidden) {{
                hiddenIds["grp_" + grp.name] = true;
                jsxLog("check group [" + grp.name + "] hidden=true");
            }} else {{
                collectHiddenIds(grp, hiddenIds);
            }}
        }} catch(e) {{}}
    }}
}}
// --- Check if layer has a pixel mask ---
function hasMask(layer) {{
    try {{
        var ref = new ActionReference();
        ref.putIdentifier(charIDToTypeID("Lyr "), layer.id);
        var desc = executeActionGet(ref);
        return desc.hasKey(stringIDToTypeID("userMaskEnabled"));
    }} catch(e) {{ return false; }}
}}
// --- Helper: apply pixel mask on the currently-active layer ---
function applyMaskAction() {{
    var desc = new ActionDescriptor();
    var ref  = new ActionReference();
    ref.putEnumerated(charIDToTypeID("Chnl"), charIDToTypeID("Chnl"), charIDToTypeID("Msk "));
    desc.putReference(charIDToTypeID("null"), ref);
    desc.putBoolean(charIDToTypeID("Aply"), true);
    executeAction(charIDToTypeID("Dlt "), desc, DialogModes.NO);
}}
// --- Apply all masks on pixel layers; non-pixel layers are handled during rasterization ---
function applyAllMasks(layerSet) {{
    for (var i = layerSet.artLayers.length - 1; i >= 0; i--) {{
        try {{
            var l = layerSet.artLayers[i];
            if (!hasMask(l)) continue;
            var lk = null;
            try {{ lk = l.kind; }} catch(e) {{}}
            var lks = "" + lk;
            var isPixel = (lk === K_NORMAL || lk === 1 || lks === "LayerKind.NORMAL");
            if (!isPixel) continue; // SOLIDFILL / shape layers — mask applied after rasterize
            app.activeDocument.activeLayer = l;
            applyMaskAction();
            jsxLog("Applied mask: " + l.name);
        }} catch(e) {{ jsxLog("mask err [" + l.name + "]: " + e.message); }}
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        try {{ applyAllMasks(layerSet.layerSets[j]); }} catch(e) {{}}
    }}
}}
// --- Step 1: Unlock ALL layers recursively (must run before delete) ---
function unlockAll(layerSet) {{
    for (var i = 0; i < layerSet.artLayers.length; i++) {{
        var l = layerSet.artLayers[i];
        try {{ if (l.isBackgroundLayer) {{ l.isBackgroundLayer = false; }} }} catch(e) {{}}
        try {{ l.allLocked = false; }} catch(e) {{ jsxLog("allLocked fail [" + l.name + "]: " + e.message); }}
        try {{ l.pixelsLocked = false; }} catch(e) {{}}
        try {{ l.positionLocked = false; }} catch(e) {{}}
        try {{ l.transparencyLocked = false; }} catch(e) {{}}
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        try {{ layerSet.layerSets[j].allLocked = false; }} catch(e) {{}}
        unlockAll(layerSet.layerSets[j]);
    }}
}}
// --- Ungroup one layer set: move all its contents to its parent, then delete it ---
function ungroupLayerSet(grp) {{
    jsxLog("ungrouping: " + grp.name);
    // Recursively dissolve any nested sub-groups first
    while (grp.layerSets.length > 0) {{
        ungroupLayerSet(grp.layerSets[0]);
    }}
    // Move each top-most layer before this group (preserves top-to-bottom order)
    while (grp.artLayers.length > 0) {{
        try {{ grp.artLayers[0].move(grp, ElementPlacement.PLACEBEFORE); }} catch(e) {{ break; }}
    }}
    // Delete the now-empty group
    try {{ grp.remove(); }} catch(e) {{ jsxLog("ungroup remove err [" + grp.name + "]: " + e.message); }}
}}
// --- Dissolve all layer groups in the document (called with workDoc) ---
function ungroupAll(layerSet) {{
    // Snapshot first: moving layers modifies the live layerSets collection
    var groups = [];
    for (var j = 0; j < layerSet.layerSets.length; j++) {{ groups.push(layerSet.layerSets[j]); }}
    for (var k = 0; k < groups.length; k++) {{
        try {{ ungroupLayerSet(groups[k]); }} catch(e) {{ jsxLog("ungroupAll err: " + e.message); }}
    }}
}}
// --- Merge all clipping mask layers into their base layer ---
// A clipping mask layer has layer.grouped === true. We merge it downward
// (into the layer immediately below) repeatedly until no grouped layers remain.
function flattenClippingMasks(doc) {{
    var changed = true;
    while (changed) {{
        changed = false;
        var layers = doc.artLayers;
        for (var i = 0; i < layers.length; i++) {{
            try {{
                if (layers[i].grouped) {{
                    jsxLog("Merging clipping mask: " + layers[i].name);
                    doc.activeLayer = layers[i];
                    layers[i].merge();
                    changed = true;
                    break; // layer collection changed — restart scan
                }}
            }} catch(e) {{ jsxLog("clipping merge err [" + layers[i].name + "]: " + e.message); break; }}
        }}
    }}
}}
// --- Step 2: Delete hidden layers using pre-collected IDs (avoids post-unlock state changes) ---
function deleteHiddenById(layerSet, hiddenIds) {{
    for (var i = layerSet.artLayers.length - 1; i >= 0; i--) {{
        try {{
            var l = layerSet.artLayers[i];
            if (hiddenIds[l.id]) {{
                jsxLog("deleting hidden: " + l.name);
                try {{ l.remove(); }} catch(e) {{ jsxLog("delete fail [" + l.name + "]: " + e.message); }}
            }}
        }} catch(e) {{ jsxLog("err delete layer: " + e.message); }}
    }}
    for (var j = layerSet.layerSets.length - 1; j >= 0; j--) {{
        try {{
            var grp = layerSet.layerSets[j];
            if (hiddenIds["grp_" + grp.name]) {{
                jsxLog("deleting hidden group: " + grp.name);
                try {{ grp.remove(); }} catch(e) {{ jsxLog("delete group fail: " + e.message); }}
            }} else {{
                deleteHiddenById(grp, hiddenIds);
            }}
        }} catch(e) {{}}
    }}
}}
// --- Check if a document contains at least one text layer ---
function hasTextLayers(layerSet) {{
    for (var i = 0; i < layerSet.artLayers.length; i++) {{
        try {{
            var lk = null;
            try {{ lk = layerSet.artLayers[i].kind; }} catch(e) {{}}
            var lks = "" + lk;
            if (lk === K_TEXT || lk === 2 || lks === "LayerKind.TEXT") return true;
        }} catch(e) {{}}
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        try {{ if (hasTextLayers(layerSet.layerSets[j])) return true; }} catch(e) {{}}
    }}
    return false;
}}
// --- Step 3: Count remaining visible layers per base name ---
function countVisible(layerSet, counts) {{
    for (var i = 0; i < layerSet.artLayers.length; i++) {{
        var b = getBaseName(layerSet.artLayers[i]);
        counts[b] = (counts[b] || 0) + 1;
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        countVisible(layerSet.layerSets[j], counts);
    }}
}}
function renamePSDLayer(layer, counters, totalCounts) {{
    var base = getBaseName(layer);
    if (!counters[base]) counters[base] = 1;
    var n = counters[base]++;
    // scenebg layers never numbered — always just "scenebg"
    if (base === NAME_PIXEL || base === NAME_SMART) return base;
    // All other types (text, stickerbg etc.): numbered only if multiple
    if (totalCounts[base] > 1) return base + n;
    return base;
}}
// --- Step 4: Rename then rasterize all remaining layers recursively ---
function renameLayers(layerSet, counters, totalCounts) {{
    for (var i = 0; i < layerSet.artLayers.length; i++) {{
        try {{
            var layer = layerSet.artLayers[i];
            var lk = null;
            try {{ lk = layer.kind; }} catch(e) {{}}
            var lks = "" + lk;
            var isText = (lk === K_TEXT || lk === 2 || lks === "LayerKind.TEXT");
            jsxLog("rename [" + layer.name + "] kind=" + lks + " isText=" + isText);
            layer.name = renamePSDLayer(layer, counters, totalCounts);
            if (!isText) {{
                try {{ layer.rasterize(RasterizeType.ENTIRELAYER); }} catch(e) {{}}
                // After rasterizing, apply any remaining mask (SOLIDFILL/shape masks)
                try {{
                    if (hasMask(layer)) {{
                        app.activeDocument.activeLayer = layer;
                        applyMaskAction();
                        jsxLog("Applied mask after rasterize: " + layer.name);
                    }}
                }} catch(me) {{ jsxLog("post-rasterize mask err [" + layer.name + "]: " + me.message); }}
            }}
        }} catch(e) {{ jsxLog("skip rename: " + e.message); }}
    }}
    for (var j = 0; j < layerSet.layerSets.length; j++) {{
        try {{ renameLayers(layerSet.layerSets[j], counters, totalCounts); }} catch(e) {{}}
    }}
}}
// --- Process one PSD file ---
function processPSD(file) {{
    var doc = null;
    var wasAlreadyOpen = false;
    // If the file is already open in Photoshop, use that instance —
    // it correctly reflects the user's saved hidden state.
    for (var d = 0; d < app.documents.length; d++) {{
        try {{
            if (app.documents[d].fullName.toString() === file.fullName.toString()) {{
                doc = app.documents[d];
                wasAlreadyOpen = true;
                jsxLog("Using already-open doc: " + file.name);
                break;
            }}
        }} catch(e) {{}}
    }}
    if (!doc) {{
        jsxLog("Opening: " + file.name);
        try {{ doc = app.open(file); }} catch(e) {{
            jsxLog("ERROR opening " + file.name + ": " + e.message);
            return;
        }}
    }}
    jsxLog("Opened: " + file.name);
    // Duplicate so we never modify the user's session document
    var workDoc = doc;
    if (wasAlreadyOpen) {{
        try {{
            workDoc = doc.duplicate(file.name, false);
            jsxLog("Duplicated for processing");
        }} catch(e) {{
            jsxLog("Dup fail: " + e.message);
            workDoc = doc;
            wasAlreadyOpen = false;
        }}
    }}
    try {{
        // STEP 1: Capture hidden state BEFORE any modifications (most accurate)
        var hiddenIds = {{}};
        collectHiddenIds(workDoc, hiddenIds);
        // STEP 2: Unlock all layers (must happen before delete to avoid error 8800)
        unlockAll(workDoc);
        jsxLog("Unlocked all layers");
        // STEP 3: Delete layers that were hidden at capture time
        deleteHiddenById(workDoc, hiddenIds);
        jsxLog("Deleted hidden layers");
        // STEP 3.5: Apply all layer masks (must run after unlockAll)
        applyAllMasks(workDoc);
        jsxLog("Applied all masks");
        // STEP 3.6: Dissolve all layer groups (ungroup)
        ungroupAll(workDoc);
        jsxLog("Ungrouped all groups");
        // STEP 3.65: Merge clipping masks into their base layers
        flattenClippingMasks(workDoc);
        jsxLog("Flattened clipping masks");
        // STEP 3.7: Skip PSD if no text layers remain after deletion
        if (!hasTextLayers(workDoc)) {{
            jsxLog("Skipped (no text layers): " + file.name);
            return;
        }}
        // Set global doc reference for isFrameStyleLayer + reset per-PSD cache
        g_workDoc   = workDoc;
        g_frameCache = {{}};
        var totalCounts = {{}};
        countVisible(workDoc, totalCounts);
        var counters = {{}};
        renameLayers(workDoc, counters, totalCounts);
        var outFile = new File(OUTPUT_FOLDER.fullName + "/" + file.name);
        jsxLog("Save target: " + outFile.fullName);
        var opts = new PhotoshopSaveOptions();
        opts.layers = true;
        opts.embedColorProfile = true;
        opts.annotations = false;
        opts.alphaChannels = true;
        opts.spotColors = true;
        try {{
            workDoc.saveAs(outFile, opts, true);
            jsxLog("Saved OK: " + outFile.fullName);
        }} catch(e) {{
            jsxLog("saveAs failed (" + e.number + "): " + e.message);
            try {{
                workDoc.saveAs(outFile, opts, false);
                jsxLog("Saved (fallback): " + outFile.fullName);
            }} catch(e2) {{
                jsxLog("Fallback save also failed: " + e2.message);
            }}
        }}
    }} finally {{
        try {{ workDoc.close(SaveOptions.DONOTSAVECHANGES); jsxLog("Closed: " + file.name); }} catch(e) {{}}
    }}
}}
// --- Main ---
var psdFiles = INPUT_FOLDER.getFiles(/\\.psd$/i);
jsxLog("PSD files found: " + (psdFiles ? psdFiles.length : 0));
if (!psdFiles || psdFiles.length === 0) {{
    jsxLog("No PSD files found in: " + INPUT_FOLDER.fullName);
    jsxLog("All done. Processed 0 file(s)");
}} else {{
    for (var k = 0; k < psdFiles.length; k++) {{
        if (STOP_FILE.exists) {{
            jsxLog("Stopped by user at file " + (k+1));
            break;
        }}
        processPSD(psdFiles[k]);
    }}
    jsxLog("All done. Processed " + psdFiles.length + " file(s)");
}}
"""
    return jsx
# ── Entry ─────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  PSD 图层重命名工具")
    print("  访问: http://127.0.0.1:7861")
    print("=" * 60)
    app.run(host="127.0.0.1", port=7861, debug=False, use_reloader=True)
